import streamlit as st
import pandas as pd
from datetime import datetime, timedelta, date 
import os
import numpy as np 
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

# --- Configuration & Data Store ---
INVENTORY_FILE = 'df.csv'
DATE_FORMAT_STRING = '%Y-%m-%d' # ISO standard for internal/form use
DATE_DISPLAY_FORMAT = 'MM/DD/YYYY' 
EXPIRY_WARNING_DAYS = 7 # Near expiry
EXPIRY_CRITICAL_DAYS = 3 # Already expired / critical warning

MASTER_COLUMNS = [
    'serial', 'segment', 'source', 'blood_type', 'component', 'volume', 
    'collected', 'expiry', 'age', 'status', 'patient'
]

# Standard Blood Types and Components for Select boxes
BLOOD_TYPES = ['O+', 'O-', 'A+', 'A-', 'B+', 'B-', 'AB+', 'AB-']
COMPONENTS = ['Whole Blood', 'PRBC', 'Platelets', 'FFP']
STATUS_OPTIONS = ['Available', 'Crossmatched', 'Expired', 'Transfused']

# --- Utility Functions (Python Reimplementation) ---

def calculate_expiry(collected_date, component):
    """Calculates the expiry date based on the component and collected date."""
    if collected_date is None:
        return None
    
    if isinstance(collected_date, datetime):
        collected_date = collected_date.date()
    
    if component in ['PRBC', 'Whole Blood']:
        return collected_date + timedelta(days=42) # Assuming 42 days for PRBC, matching the common 6-week shelf life
    elif component == 'Platelets':
        return collected_date + timedelta(days=5)
    elif component == 'FFP':
        # HTML logic used 7 years. Let's use 7 * 365 + 1 day for leap year buffer.
        return collected_date + timedelta(days=7 * 365 + 1)
    else:
        # Default fallback, e.g., for unknown components
        return collected_date + timedelta(days=42)

def compute_age_text(collected_date, component):
    """Computes the age of the unit, formatted as 'Nd' or 'Ny Md' for FFP."""
    if collected_date is None:
        return ''
    
    if isinstance(collected_date, datetime):
        collected_date = collected_date.date()
        
    today = datetime.today().date()
    
    if collected_date > today:
        return 'Future'
        
    diff_days = (today - collected_date).days
    
    if component == "FFP":
        # Match HTML logic: display in years and remaining days
        y = diff_days // 365
        d = diff_days % 365
        return f"{y}y {d}d"
    
    return f"{diff_days}d"

def load_data():
    """Loads or initializes the inventory DataFrame."""
    if os.path.exists(INVENTORY_FILE):
        df = pd.read_csv(INVENTORY_FILE)
    else:
        df = pd.DataFrame(columns=MASTER_COLUMNS)
    
    # CRITICAL: Ensure date columns are proper datetime objects
    for col in ['collected', 'expiry']:
        if col in df.columns:
            # Replace 'None' string with NaN for clean conversion
            df[col] = df[col].replace('None', np.nan) 
            df[col] = pd.to_datetime(df[col], format=DATE_FORMAT_STRING, errors='coerce')
        else:
            df[col] = pd.NaT 

    # Fill NaNs in text columns
    for col in ['segment', 'source', 'patient']:
        if col in df.columns:
            df[col] = df[col].fillna('None').astype(str)
            
    # Re-calculate age and status on load
    df['age'] = df.apply(lambda row: compute_age_text(row['collected'], row['component']), axis=1)
    df = update_inventory_status(df)
        
    df = df.reindex(columns=MASTER_COLUMNS, fill_value=None)
    return df

def update_inventory_status(df):
    """Checks expiry dates and updates unit status."""
    today_date_only = datetime.today().date()
    expirable_statuses = ['Available', 'Crossmatched']
    
    # Check for expired units
    if 'expiry' in df.columns:
        # Check if the date is less than or equal to today, and if it's one of the expirable statuses
        # The .dt.date access is now safe because load_data enforces datetime type.
        df.loc[(df['status'].isin(expirable_statuses)) & (df['expiry'].dt.date <= today_date_only), 'status'] = 'Expired'
    
    return df

def save_data(df):
    """Saves the current DataFrame back to the CSV file using ISO format."""
    
    def date_to_string(d):
        if pd.isna(d): 
            return 'None'
        try:
            # Save using standard ISO format YYYY-MM-DD
            return d.strftime(DATE_FORMAT_STRING)
        except:
            return 'None'
    
    # Save a copy with dates as ISO strings
    df_save = df.copy()
    df_save['collected'] = df_save['collected'].apply(date_to_string)
    df_save['expiry'] = df_save['expiry'].apply(date_to_string)
    
    df_save.to_csv(INVENTORY_FILE, index=False)

def color_rows_by_expiry(row):
    """Apply CSS styling based on proximity to expiry date."""
    if row['status'] == 'Expired':
        # already-expired equivalent: light red background
        return ['background-color: #f8d7da'] * len(row) 
        
    if pd.isna(row['expiry']):
        return [''] * len(row)

    days_left = (row['expiry'].date() - datetime.today().date()).days
    
    if days_left <= EXPIRY_CRITICAL_DAYS and days_left > 0:
        # near-expiry equivalent: light yellow background
        return ['background-color: #fff3cd'] * len(row)
    
    return [''] * len(row)

# --- Report Generation (Python-Docx implementation) ---

def generate_docx_report(df_active):
    """Generates a DOCX file based on active inventory, grouped by component and blood type."""
    doc = Document()
    doc.add_heading('Daily Blood Inventory Report', 0)
    doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    
    df_active = df_active[df_active['status'].isin(['Available', 'Crossmatched'])].copy()
    
    # Data structure: Component -> Blood Type -> List of units
    inventory_data = {comp: {bt: [] for bt in BLOOD_TYPES} for comp in COMPONENTS}
    
    for index, row in df_active.iterrows():
        comp = row['component']
        bt = row['blood_type']
        
        if comp in inventory_data and bt in inventory_data[comp]:
            # Calculate display string: expiry for FFP, age for others (matching JS logic)
            age_or_expiry = ''
            if comp == 'FFP':
                if pd.notna(row['expiry']):
                    age_or_expiry = row['expiry'].strftime('%b %d, %Y')
            else:
                age_or_expiry = compute_age_text(row['collected'], comp)

            inventory_data[comp][bt].append({
                'serial': row['serial'],
                'ageOrExpiry': age_or_expiry,
                'patient': row['patient'] if row['patient'] != 'None' else ''
            })

    # Add a table for each component
    for component_name in COMPONENTS:
        component_map = inventory_data[component_name]
        
        doc.add_heading(component_name, level=2)
        
        # Determine max rows for this component's table
        max_rows = max([len(units) for units in component_map.values()], default=0)
        max_rows = max(max_rows, 1) # at least 1 data row
        
        # Create table with 8 columns (for the 8 blood types)
        table = doc.add_table(rows=max_rows + 1, cols=len(BLOOD_TYPES))
        table.style = 'Table Grid'
        
        # Header Row
        header_cells = table.rows[0].cells
        for i, bt in enumerate(BLOOD_TYPES):
            header_cells[i].text = bt
            header_cells[i].paragraphs[0].runs[0].font.bold = True
            header_cells[i].width = Inches(0.8)

        # Data Rows
        for r_index in range(max_rows):
            row_cells = table.rows[r_index + 1].cells
            for c_index, bt in enumerate(BLOOD_TYPES):
                units = component_map[bt]
                if r_index < len(units):
                    item = units[r_index]
                    p = row_cells[c_index].paragraphs[0]
                    p.add_run(f"{item['serial']} — {item['ageOrExpiry']}").font.size = Pt(9)
                    if item['patient']:
                        p.add_run(f"\nPatient: {item['patient']}").italic = True
                        p.runs[-1].font.size = Pt(8)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Save to a BytesIO object
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- Application Initialization ---

if 'inventory_df' not in st.session_state:
    st.session_state['inventory_df'] = load_data()

# Ensure status is updated at every rerun (for expiry checks)
st.session_state['inventory_df'] = update_inventory_status(st.session_state['inventory_df'].copy())

# --- Streamlit UI ---

st.set_page_config(layout="wide", page_title="Blood Bag Inventory")

# --- Header & Clock (Replaces HTML top section) ---
st.markdown("""
<style>
    /* Mimic HTML floating/fixed element for the clock and status */
    #liveClock {
        font-weight: 600;
        text-align: right;
    }
    .top-header {
        display: flex;
        justify-content: space-between;
        align-items: center;
        margin-bottom: 1rem;
    }
</style>
""", unsafe_allow_html=True)

# Using st.empty() for the live clock update
col_title, col_clock = st.columns([3, 1])

with col_title:
    st.markdown("<h1 class='h3 mb-0'>Blood Bag Inventory</h1>", unsafe_allow_html=True)
    st.markdown("<small class='text-muted'>Manage units • Live updates</small>", unsafe_allow_html=True)

# We cannot implement a true live clock in Streamlit due to its stateless nature. 
# We'll display the current time on load.
with col_clock:
    # Use an empty placeholder for a mock 'live' clock (updates on rerun)
    clock_placeholder = st.empty()
    clock_placeholder.markdown(f"**⏱️ {datetime.now().strftime('%b %d, %Y — %H:%M:%S')}**")
    
st.markdown("---")


# --- Sidebar Global Filters (Replaces HTML floatingFilters) ---
st.sidebar.header("🔍 Search & Filters")

q_search = st.sidebar.text_input("Search (Serial / Blood / Patient)")
f_blood = st.sidebar.selectbox("Blood Type", ['All'] + BLOOD_TYPES)
f_component = st.sidebar.selectbox("Component", ['All'] + COMPONENTS)
f_status = st.sidebar.selectbox("Status", ['All'] + STATUS_OPTIONS)

# --- Data Filtering Logic ---

filtered_df = st.session_state['inventory_df'].copy()

# Apply text search
if q_search:
    q_search = q_search.lower()
    filtered_df = filtered_df[
        filtered_df['serial'].str.lower().str.contains(q_search, na=False) |
        filtered_df['blood_type'].str.lower().str.contains(q_search, na=False) |
        filtered_df['patient'].str.lower().str.contains(q_search, na=False)
    ]

# Apply selectbox filters
if f_blood != 'All':
    filtered_df = filtered_df[filtered_df['blood_type'] == f_blood]
if f_component != 'All':
    filtered_df = filtered_df[filtered_df['component'] == f_component]

# --- Main Tabs (Replaces HTML nav-pills) ---
tab_add, tab_inv, tab_exp, tab_trans = st.tabs(["➕ Add Blood Bag", "🔬 Inventory", "🔴 Expired", "💉 Transfused"])


# ==================================
# TAB 1: ADD BLOOD BAG
# ==================================
with tab_add:
    st.header("Add New Blood Bag")
    
    # Use st.form for atomic submission (replaces HTML form)
    with st.form("new_unit_form", clear_on_submit=True):
        st.subheader("Unit Details")
        col_id, col_seg, col_src = st.columns(3)
        serial = col_id.text_input("Serial Number", key='add_serial', required=True)
        segment = col_seg.text_input("Segment Number", key='add_segment', required=True)
        source = col_src.text_input("Source", key='add_source', value='Donor')

        col_group, col_comp, col_vol, col_coll = st.columns(4)
        blood_type = col_group.selectbox("Blood Type", BLOOD_TYPES, key='add_blood_type', required=True)
        
        # Dynamic Expiry Calculation based on Component
        component = col_comp.selectbox("Component", COMPONENTS, key='add_component', required=True)
        volume = col_vol.number_input("Volume (mL)", min_value=1, value=450, step=10, key='add_volume', required=True)
        collected_date_input = col_coll.date_input("Collection Date", value='today', key='add_collected', required=True)
        
        # Calculate Expiry Date (replicates JS calculation logic)
        calculated_expiry = calculate_expiry(collected_date_input, component)
        
        col_exp, col_status, col_pat = st.columns(3)
        expiry_date_input = col_exp.date_input(
            "Expiry Date (Auto-calculated, Edit if FFP)", 
            value=calculated_expiry, 
            key='add_expiry'
        )
        
        initial_status = col_status.selectbox("Status", STATUS_OPTIONS, key='add_status')
        patient = col_pat.text_input("Patient Name (Required if Crossmatched/Transfused)", key='add_patient', value='')

        st.markdown("---")
        submit_button = st.form_submit_button("➕ Add Blood Bag", type="primary")
        
        if submit_button:
            if serial in st.session_state['inventory_df']['serial'].values:
                st.error(f"Unit Serial Number {serial} already exists.")
            else:
                new_unit_data = {
                    'serial': serial,
                    'segment': segment,
                    'source': source,
                    'blood_type': blood_type,
                    'component': component,
                    'volume': volume,
                    'collected': pd.to_datetime(collected_date_input),
                    'expiry': pd.to_datetime(expiry_date_input),
                    'age': compute_age_text(collected_date_input, component),
                    'status': initial_status,
                    'patient': patient if patient else 'None'
                }

                new_unit_df = pd.DataFrame([new_unit_data], columns=MASTER_COLUMNS)
                
                # Use load_data to get the string-formatted DataFrame for saving
                final_concatenated_df = pd.concat([st.session_state['inventory_df'], new_unit_df], ignore_index=True)
                
                save_data(final_concatenated_df)
                
                # Re-load the session state with the new data as datetime objects
                st.session_state['inventory_df'] = load_data() 
                st.success(f"Unit {serial} successfully added! Status: {initial_status}")
                st.rerun() # Refresh the view
            
# ==================================
# TAB 2: ACTIVE INVENTORY
# ==================================
with tab_inv:
    st.header("Active Inventory")
    
    # Filter for active units only, then apply sidebar filters
    inventory_df_filtered = filtered_df[filtered_df['status'].isin(['Available', 'Crossmatched'])].copy()
    
    # Sorting (matches HTML logic: status False/Expired first, then Expiry True/Soonest first)
    inventory_df_filtered = inventory_df_filtered.sort_values(
        by=['status', 'expiry'], 
        ascending=[True, True]
    )

    if inventory_df_filtered.empty:
        st.info("No active units found matching the current filters.")
    else:
        # Renaming columns for display clarity (matching HTML table headers)
        display_df = inventory_df_filtered.rename(columns={
            'serial': 'Serial', 'segment': 'Segment', 'source': 'Source', 
            'blood_type': 'Blood Type', 'component': 'Component', 
            'volume': 'Volume', 'collected': 'Collection', 'expiry': 'Expiry', 
            'age': 'Days Old', 'status': 'Status', 'patient': 'Patient'
        })[['Serial', 'Segment', 'Source', 'Blood Type', 'Component', 'Volume', 'Collection', 'Expiry', 'Days Old', 'Status', 'Patient']]
        
        st.caption("Editing the table below automatically updates the inventory data.")

        # Use st.data_editor to allow inline editing (replaces HTML table editing)
        edited_df = st.data_editor(
            display_df.style.apply(color_rows_by_expiry, axis=1),
            key="active_inventory_editor",
            use_container_width=True,
            column_config={
                "Collection": st.column_config.DateColumn("Collection", format=DATE_DISPLAY_FORMAT, disabled=True),
                "Expiry": st.column_config.DateColumn("Expiry", format=DATE_DISPLAY_FORMAT),
                "Status": st.column_config.SelectboxColumn("Status", options=STATUS_OPTIONS),
            },
            num_rows="fixed" # prevent accidental deletion, use action column if needed
        )
        
        # Handle saving edits from the data_editor
        if st.button("💾 Save Changes", type="primary"):
            # Update the original dataframe in session state
            for col in ['Collection', 'Expiry']:
                edited_df[col] = pd.to_datetime(edited_df[col], errors='coerce')
                
            edited_df.columns = [c.lower().replace(' ', '_') for c in edited_df.columns]
            
            # Match edited rows back to the session state dataframe
            st.session_state['inventory_df'].loc[edited_df.index] = edited_df.values
            st.session_state['inventory_df'] = update_inventory_status(st.session_state['inventory_df'])
            save_data(st.session_state['inventory_df'])
            st.success("Inventory changes saved and updated!")
            st.rerun()

        st.markdown("---")
        
        # Daily Report Button (Replaces HTML dailyReportBtn)
        report_bytes = generate_docx_report(inventory_df_filtered)
        st.download_button(
            label="Daily Report (DOCX)",
            data=report_bytes,
            file_name=f"Daily_Blood_Inventory_Report_{datetime.now().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type='secondary'
        )

# ==================================
# TAB 3: EXPIRED UNITS
# ==================================
with tab_exp:
    st.header("Expired Units")
    expired_df = filtered_df[filtered_df['status'] == 'Expired'].copy()
    
    if expired_df.empty:
        st.info("No expired units found matching the current filters.")
    else:
        st.dataframe(expired_df.style.apply(color_rows_by_expiry, axis=1), use_container_width=True)

# ==================================
# TAB 4: TRANSFUSED UNITS
# ==================================
with tab_trans:
    st.header("Transfused Units")
    transfused_df = filtered_df[filtered_df['status'] == 'Transfused'].copy()
    
    if transfused_df.empty:
        st.info("No transfused units found matching the current filters.")
    else:
        st.dataframe(transfused_df, use_container_width=True)
